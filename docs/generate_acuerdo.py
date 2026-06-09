"""Generate DOCX for ACUERDO_ACCESO_OPENACM."""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE = RGBColor(0x1A, 0x37, 0x6C)
DARK_GRAY = RGBColor(0x55, 0x55, 0x55)

MD_PATH = os.path.join(os.path.dirname(__file__), "ACUERDO_ACCESO_OPENACM.md")
DOCX_PATH = os.path.join(os.path.dirname(__file__), "ACUERDO_ACCESO_OPENACM.docx")


def sp(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)
    sp(p, 2, 2)


def add_left_border(para, color="1A376C"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def inline(text, para, size=11, bold=False, color=None):
    """Add runs to para, handling **bold** markers."""
    for seg in re.split(r"(\*\*[^*]+\*\*)", text):
        is_bold = seg.startswith("**") and seg.endswith("**")
        content = seg[2:-2] if is_bold else seg
        if not content:
            continue
        r = para.add_run(content)
        r.bold = is_bold or bold
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        if color:
            r.font.color.rgb = color


def build_docx():
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

    # Normal style baseline
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    with open(MD_PATH, encoding="utf-8") as f:
        lines = [l.rstrip("\n") for l in f.readlines()]

    i = 0
    in_code = False

    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        # Code fence toggle
        if s.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        # Monospace block (signature area)
        if in_code:
            if s:
                p = doc.add_paragraph()
                r = p.add_run(s)
                r.font.name = "Courier New"
                r.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(1.5)
                sp(p, 0, 1)
            i += 1
            continue

        # Skip empty lines
        if not s:
            i += 1
            continue

        # H1 — document title
        if s.startswith("# ") and not s.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(s[2:])
            r.bold = True
            r.font.size = Pt(16)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK_BLUE
            sp(p, 0, 14)
            i += 1
            continue

        # H2
        if s.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(s[3:].upper())
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK_BLUE
            sp(p, 12, 4)
            i += 1
            continue

        # H3 — clause title
        if s.startswith("### "):
            p = doc.add_paragraph()
            inline(s[4:], p, size=11, bold=True, color=DARK_BLUE)
            sp(p, 10, 3)
            i += 1
            continue

        # Horizontal rule
        if s == "---":
            add_hr(doc)
            i += 1
            continue

        # Blockquote — collect consecutive lines
        if s.startswith("> "):
            bq = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                bq.append(lines[i].strip()[2:])
                i += 1
            for line in bq:
                if not line.strip():
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.7)
                p.paragraph_format.right_indent = Cm(0.5)
                add_left_border(p)
                inline(line, p, size=10, color=DARK_GRAY)
                sp(p, 1, 2)
            continue

        # Bullet list
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            inline(s[2:], p, size=11)
            sp(p, 0, 3)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            inline(m.group(1), p, size=11)
            sp(p, 0, 3)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        inline(s, p, size=11)
        sp(p, 0, 5)
        i += 1

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
    print("Done.")
